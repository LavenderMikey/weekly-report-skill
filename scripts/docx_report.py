# -*- coding: utf-8 -*-
"""Report-agnostic helpers for building a weekly report by CLONING the user's
historical .docx and re-emitting content from harvested prototype paragraphs.

These are the reusable building blocks the SKILL.md "clone-and-rebuild-from-
prototypes" recipe refers to. They are generic — nothing here is specific to a
particular report. A typical build looks like:

    import shutil, docx
    from docx_report import (wipe_body, emit, clone_with_text,
                             clone_with_omath, add_figure_table,
                             strip_unused_media)
    from docx_math import latex_to_omath
    from copy import deepcopy

    shutil.copy(template_path, out_path)
    doc = docx.Document(out_path)
    body = doc.element.body
    paras = body.findall("{...}p")           # qn('w:p')

    # 1) harvest prototypes from the most recent week (indices come from
    #    inspect_docx.py output: a top-level numbered heading, a subsection
    #    numbered heading, a body paragraph, a plan item, the title line, ...)
    P_SECTION = deepcopy(paras[i_section])   # numId for 综述/papers -> 一、二、三
    P_SUB     = deepcopy(paras[i_sub])       # subsection numId     -> 1、2、3
    P_BODY    = deepcopy(paras[i_body])      # numId=0 / no number
    ...

    # 2) wipe and rebuild
    sectPr = wipe_body(doc)
    emit(doc, sectPr, clone_with_text(P_SECTION, "综述"))
    emit(doc, sectPr, clone_with_text(P_BODY, "本周阅读了…"))
    emit(doc, sectPr, clone_with_omath(P_BODY, latex_to_omath([r"E=mc^2"])[0]))
    add_figure_table(doc, sectPr, "fig.png", "图 1  整体架构")
    ...
    doc.save(out_path)
    strip_unused_media(out_path)   # drop old weeks' images so the file stays small

Why clone? A real report inherits its fonts, line/paragraph spacing, automatic
numbering (numbering.xml) and page setup from styles the text can't show.
Cloning keeps all of that; cloned headings that reuse the same numId get
renumbered by Word automatically.
"""
import os
import shutil
import zipfile

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _w(tag):
    return qn("w:" + tag)


def wipe_body(doc):
    """Remove ALL body content (paragraphs and tables left from prior weeks)
    except the trailing <w:sectPr> (page setup). Returns the sectPr element so
    new content can be inserted before it. Harvest your prototypes BEFORE
    calling this."""
    body = doc.element.body
    sectPr = body.find(_w("sectPr"))
    for child in list(body):
        if child.tag != _w("sectPr"):
            body.remove(child)
    return sectPr


def emit(doc, sectPr, el):
    """Insert a block element (paragraph/table) at the current end of content,
    i.e. just before the section properties, preserving document order."""
    if sectPr is not None:
        sectPr.addprevious(el)
    else:
        doc.element.body.append(el)


def clone_with_text(proto, text):
    """Deep-copy a prototype <w:p>, keep its first run's formatting and its pPr
    (including numPr, so numbering survives), and replace the text. Extra runs
    and any inline images on the prototype are dropped."""
    from copy import deepcopy
    el = deepcopy(proto)
    runs = el.findall(_w("r"))
    if runs:
        first = runs[0]
        for r in runs[1:]:
            el.remove(r)
        for child in list(first):
            if child.tag in (_w("t"), _w("drawing"), _w("tab"), _w("br"),
                             _w("pict")):
                first.remove(child)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        first.append(t)
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        el.append(r)
    return el


def clone_with_omath(proto, omath_el):
    """Deep-copy a prototype <w:p>, strip its runs/equations, and append a Word
    OMML element (from docx_math.latex_to_omath). Inherits the prototype's
    paragraph formatting (e.g. indent), with no numbering if the prototype was a
    body paragraph."""
    from copy import deepcopy
    el = deepcopy(proto)
    for r in el.findall(_w("r")):
        el.remove(r)
    for tag in ("oMath", "oMathPara"):
        for e in el.findall("{%s}%s" % (M_NS, tag)):
            el.remove(e)
    el.append(omath_el)
    return el


def set_table_borders(tbl):
    """Single-line borders all around + between cells, without depending on a
    'Table Grid' style that the cloned template may not define."""
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)


def add_figure_table(doc, sectPr, png, caption, width_cm=14.0, border=True):
    """Insert a centered 2-row table (image row + caption row) at the current
    end of content. 'Inserting the image via a table', positioned in document
    order — call this right after the body paragraph the figure belongs under,
    so it lands inline like the template (not bunched at the section end)."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if border:
        set_table_borders(tbl)
    cp = tbl.cell(0, 0).paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.add_run().add_picture(png, width=Cm(width_cm))
    capp = tbl.cell(1, 0).paragraphs[0]
    capp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    capp.add_run(caption)
    # add_table appends to the body end (after sectPr); move it into position
    if sectPr is not None:
        sectPr.addprevious(tbl._tbl)
    return tbl


def add_centered_image(doc, sectPr, proto_body, png, width_cm=14.0):
    """Alternative to a table: insert the image as a centered paragraph cloned
    from a body prototype — matches templates that place figures as plain
    centered inline images (no border). Returns the paragraph element."""
    from copy import deepcopy
    el = deepcopy(proto_body)
    for r in el.findall(_w("r")):
        el.remove(r)
    emit(doc, sectPr, el)
    # wrap as python-docx Paragraph to use add_run().add_picture
    from docx.text.paragraph import Paragraph
    p = Paragraph(el, doc.element.body)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(png, width=Cm(width_cm))
    return el


def strip_unused_media(path):
    """Cloning keeps every image from the old multi-week file in word/media/,
    even after their paragraphs are removed. Drop media not referenced by
    document.xml (and their relationship entries) so the file stays small.
    Returns the number of images removed."""
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        data = {n: z.read(n) for n in names}
    droot = etree.fromstring(data["word/document.xml"])
    used = set()
    for el in droot.iter():
        for k, v in el.attrib.items():
            if k.startswith("{%s}" % R):
                used.add(v)
    rroot = etree.fromstring(data["word/_rels/document.xml.rels"])
    drop_files = set()
    for rel in list(rroot):
        rid = rel.get("Id")
        target = rel.get("Target", "")
        if "media/" in target and rid not in used:
            rroot.remove(rel)
            drop_files.add("word/" + target.lstrip("/"))
    new_rels = etree.tostring(rroot, xml_declaration=True,
                              encoding="UTF-8", standalone=True)
    tmp = path + ".tmp"
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for n in names:
            if n in drop_files:
                continue
            z.writestr(n, new_rels if n == "word/_rels/document.xml.rels"
                       else data[n])
    shutil.move(tmp, path)
    return len(drop_files)
